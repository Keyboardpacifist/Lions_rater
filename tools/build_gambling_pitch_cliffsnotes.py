"""Cliff's Notes version of the gambling product pitch — same ideas
as the full doc, ~20% of the words. Written for a sharp who already
knows the gambling space and wants the punchlines.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


REPO = Path(__file__).resolve().parent.parent
OUT = REPO / "Lions_Rater_Gambling_CliffsNotes.docx"


def _h1(doc, text):
    return doc.add_heading(text, level=1)


def _h2(doc, text):
    return doc.add_heading(text, level=2)


def _p(doc, text):
    return doc.add_paragraph(text)


def _bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def _bold_label(doc, label, body):
    p = doc.add_paragraph()
    run = p.add_run(label + " ")
    run.bold = True
    p.add_run(body)
    return p


def build():
    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    # ── Title ────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Lions Rater — Gambling Product Vision (Cliff's Notes)"
    )
    run.bold = True
    run.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(
        "The 1-page version. Full deck available."
    )
    sr.italic = True
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(80, 80, 90)
    doc.add_paragraph()

    thesis = doc.add_paragraph()
    thesis.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = thesis.add_run(
        "Competitors give you the who and what. We give you that plus "
        "what it means for Sunday's matchup."
    )
    tr.italic = True
    tr.bold = True
    tr.font.size = Pt(12)
    doc.add_paragraph()

    # ── The Product ──────────────────────────────────────────────
    _h1(doc, "The Product")

    _p(doc,
        "Lions Rater is a 3-year-old NFL/college player-rating app "
        "in production. We've accidentally accumulated a "
        "cross-indexed football data foundation — every play since "
        "2016 tagged by formation, personnel, weather, score, "
        "teammate availability, and coaching staff — that nobody "
        "else in the public market has. The gambling product is "
        "what gets built on top of it."
    )
    _p(doc,
        "Structural edge: we compute granular correlations books "
        "treat as independent. Same-game props, same-team usage "
        "shifts, weather × player splits, multi-injury cascading. "
        "Off-the-shelf data feeds can't match it."
    )

    # ── Market ───────────────────────────────────────────────────
    _h1(doc, "The Market")

    _bold_label(doc, "Target:",
        "Serious recreational bettors ($500–2K/week, ~2–4M Americans). "
        "Big enough to fund a single founder; too small for ESPN. "
        "Pay $30–100/mo for real edge.")

    _bold_label(doc, "Why this segment:",
        "They want transparency, education, and tools — not picks "
        "services. They're underserved because most paid tools "
        "assume sharp expertise.")

    _bold_label(doc, "Brand pillar — Bet School Mode:",
        "Two UI modes — Standard for power users, Bet School for "
        "newcomers. Same data, two reading levels. Plain-English "
        "labels, hover tooltips, auto-explanations, free Practice "
        "Mode with a fake $1K bankroll. \"Smart enough for sharps, "
        "friendly enough for fans.\"")

    # ── Pricing ──────────────────────────────────────────────────
    _h1(doc, "Pricing")

    table = doc.add_table(rows=4, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Tier"
    hdr[1].text = "Price"
    hdr[2].text = "Active"
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    rows = [
        ("Free (rater app + basics)", "$0", "Year-round"),
        ("Standard Paid", "$15–20/mo", "Sept–Feb only"),
        ("Pro (API + early alerts)", "$50–80/mo", "Sept–Feb only"),
    ]
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val

    _p(doc, "")
    _p(doc,
        "Break-even at <100 paid subs. 1,000 paid = $100K/yr. "
        "3,000 paid = $375K/yr. Reach is a primary goal alongside "
        "revenue."
    )

    # ── Costs ────────────────────────────────────────────────────
    _h1(doc, "Costs")
    _bullet(doc,
        "Live odds API (in-season): $30–60/mo × 5 mo")
    _bullet(doc,
        "Historical odds backfill (one-time): $500–1,500")
    _bullet(doc,
        "News + email + push: ~$50–80/mo total")
    _bullet(doc,
        "All player and team data: free (nflverse, CFBD, "
        "nflreadpy)")

    _p(doc,
        "Year 1 lean: $1,200–2,000. Premium with Twitter API: "
        "$2,500–4,000."
    )

    # ── Features ─────────────────────────────────────────────────
    _h1(doc, "The 6 Game-Bet Features")

    _bullet(doc,
        "1. Injury Probability + Usage Retention — historical "
        "cohort model with named comparable cases.")
    _bullet(doc,
        "2. Game-Script Simulator — toggle starters in/out, watch "
        "pass rate / personnel / props re-price live; handles "
        "cascading multi-injury.")
    _bullet(doc,
        "3. Books-vs-Model Behavioral Baseline — historical archive "
        "of how books over- or under-react to specific kinds of "
        "injury news; surfaces persistent edge.")
    _bullet(doc,
        "4. Smart Alerts + Digests — every alert fuses news, "
        "probability, game-script, prop re-pricing, and personal "
        "bet impact into one push.")
    _bullet(doc,
        "5. Weather Production Window — slider-driven weather "
        "modeling with player-cohort-based P10/P50/P90 production "
        "ranges; sample-size confidence flags.")
    _bullet(doc,
        "6. Coaching & Scheme Tendency Layer — career play-caller "
        "tendencies (pass/run, RZ aggression, play-action, blitz "
        "rate) feeding team-level projections, traveling with the "
        "coach when he changes jobs.")

    _h1(doc, "The 9 Prop-Bet Features")

    _bullet(doc,
        "1. Decomposed Prop Projection — every projection ships "
        "with a line-by-line breakdown of where each yard came "
        "from. Auditable, not black-box.")
    _bullet(doc,
        "2. SGP Correlation Edge — books price same-game parlays "
        "as independent. They aren't. We compute real correlation "
        "from pbp and surface mispriced SGPs.")
    _bullet(doc,
        "3. Alt-Line EV Finder (with Ladder for Dummies) — across "
        "the ladder, find the rung where the book is most wrong. "
        "Bet School Mode teaches alt-lines along the way.")
    _bullet(doc,
        "4. Smart Parlay Builder — daily mispriced-leg menu, "
        "correlation-aware auto-suggestions, conflict detection, "
        "auto-optimized parlays by EV.")
    _bullet(doc,
        "5. Anytime / First TD Probability Vector — RZ snap share "
        "+ goal-line defense + game script → per-player TD "
        "probability split by rush/rec.")
    _bullet(doc,
        "6. Snap-Share / Target-Share Trend Divergence — flags "
        "props where last-3-week usage decoupled from season "
        "baseline. Books anchor to season; we catch the shift.")
    _bullet(doc,
        "7. Longest-Play Edge Finder — explosive-rate z-scores find "
        "mispriced longest-rush / longest-reception props.")
    _bullet(doc,
        "8. Defense-vs-Position Stat Allowance Tables — "
        "foundational matchup research; sortable per defense per "
        "stat per opposing position.")
    _bullet(doc,
        "9. Coaching Tendency Adjustment — OC/DC tendency-based "
        "per-player adjustments. Screen-heavy OC inflates RB/slot "
        "props; disguise-heavy DC inflates QB sack/INT props.")

    # ── Edge moat ────────────────────────────────────────────────
    _h1(doc, "Why the Edge Persists at Scale")

    _p(doc,
        "Most gambling tools have the picks-service trap: tell "
        "10K subs the same lean, books move the line, edge dies. "
        "Our product avoids it because the scenario builder is a "
        "research environment, not a picks service. Two subscribers "
        "stare at the same scenario tree and place different bets "
        "based on their own injury read, their own correlation "
        "tolerance, their own parlay preferences. The downstream "
        "bets diverge because human judgment is the last mile. "
        "Edge persists even at 10K+ subscribers."
    )

    # ── Questions ────────────────────────────────────────────────
    _h1(doc, "Questions for the Sharp")

    _bullet(doc,
        "Which features hit hardest? Where would you actually pay?")
    _bullet(doc,
        "What's missing from this stack that you pay for elsewhere?")
    _bullet(doc,
        "Patterns of book mis-reaction we haven't listed?")
    _bullet(doc,
        "If we had to cut 3 features for v1, which 3?")
    _bullet(doc,
        "Honest read on $15–20/mo Standard, $50–80/mo Pro?")
    _bullet(doc,
        "Bet School Mode — dilution or TAM expander?")

    doc.save(OUT)
    print(f"✓ wrote {OUT}")


if __name__ == "__main__":
    build()
