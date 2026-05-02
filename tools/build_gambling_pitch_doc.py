"""Generate the full Word doc summarizing the gambling-product vision
for the Lions Rater app. Written for a meeting with a sharp from
Vegas; tone is concrete and honest about constraints.

Locked feature count:
  - 6 main / game-bet features
  - 9 prop-bet features
  - Brand pillar: Standard Mode vs Bet School Mode
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


REPO = Path(__file__).resolve().parent.parent
OUT = REPO / "Lions_Rater_Gambling_Product_Vision.docx"


def _h1(doc, text):
    return doc.add_heading(text, level=1)


def _h2(doc, text):
    return doc.add_heading(text, level=2)


def _h3(doc, text):
    return doc.add_heading(text, level=3)


def _p(doc, text):
    return doc.add_paragraph(text)


def _bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def _bold_label(doc, label, body):
    p = doc.add_paragraph()
    run = p.add_run(label + " ")
    run.bold = True
    p.add_run(body)
    return p


def _italic_quote(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.bold = True
    run.font.size = Pt(13)
    return p


def _feature_block(doc, title, what, how, sample, diff):
    _h3(doc, title)
    _bold_label(doc, "What it does.", what)
    _bold_label(doc, "How it works.", how)
    _bold_label(doc, "Sample output.", sample)
    _bold_label(doc, "Differentiation.", diff)


def build():
    doc = Document()

    # Margins
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    # ── Title ────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Lions Rater — Sports Gambling Product Vision")
    run.bold = True
    run.font.size = Pt(22)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(
        "A 6-feature game-bet stack and 9-feature prop-bet suite "
        "built on a cross-indexed football data foundation."
    )
    sr.italic = True
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(80, 80, 90)
    doc.add_paragraph()

    _italic_quote(doc,
        "Our competitors give you the who and what. We give you that "
        "plus what it means for Sunday's matchup."
    )

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════
    # 1. INTRODUCTION
    # ════════════════════════════════════════════════════════════
    _h1(doc, "1. Introduction — What We've Built")

    _p(doc,
        "Lions Rater is a production NFL and college football "
        "player-rating platform serving fans at "
        "lions-rater.streamlit.app. Over three years of iterative "
        "development, the app has accumulated something less obvious "
        "than a feature set: a deep, cross-indexed data foundation "
        "that — almost by accident — turns out to be uniquely suited "
        "for sports gambling research."
    )

    _p(doc,
        "The core asset is not any single dataset. It's a data bank "
        "with cross-sorting: every play, every player, every game "
        "from 2016 onward is indexed against multiple dimensions "
        "simultaneously. The data foundation pulls from nflverse "
        "play-by-play (~485,000 plays), College Football Data API, "
        "NFL Combine and Pro Day measurables, NextGenStats player "
        "tracking, Pro-Football-Reference advanced stats, weekly snap "
        "counts, draft history, OverTheCap contract data, weather "
        "observations, and 700+ college recruiting classes. Every "
        "play in our database carries the full context of when, "
        "where, and how it happened."
    )

    _p(doc,
        "The practical consequence is that any player can be queried "
        "by any combination of:"
    )
    _bullet(doc, "Position group and role within scheme")
    _bullet(doc, "Formation and personnel grouping (11, 12, 21, etc.)")
    _bullet(doc,
        "Down, distance, yard line, quarter, score differential")
    _bullet(doc, "Weather conditions, stadium, and field surface")
    _bullet(doc,
        "Game-script context (winning, trailing, blowout, garbage time)")
    _bullet(doc,
        "Teammate availability — which starters were active or limited")
    _bullet(doc,
        "Defensive opponent tendencies and personnel matchups")
    _bullet(doc, "Season, week, matchup history, and recency")
    _bullet(doc, "Coaching staff and scheme tendencies")

    _p(doc,
        "Most public gambling-research products are built on a single "
        "dataset and answer a single dimension at a time. They can "
        "tell you Joe Burrow's career pass yards in cold weather. "
        "They cannot tell you Burrow's pass yards in cold weather, "
        "against Cover-3 zone, when the Bengals are trailing, with "
        "Tee Higgins on the injury report. We can. For every player. "
        "For every play. Since 2016."
    )

    _p(doc,
        "The most valuable consequence is correlation. Sportsbooks "
        "price most multi-leg bets as if their legs are independent. "
        "They aren't. When Goff goes over his pass yards, the "
        "empirical conditional probability that St. Brown also goes "
        "over his receiving yards isn't 50% — it's typically 70–80%, "
        "because both events flow from the same plays. We can compute "
        "these correlations across every player pair, every team "
        "total, every prop combination, going back nearly a decade. "
        "That ability — to model granular conditional probabilities "
        "books treat as independent — is the structural edge our "
        "product is built on."
    )

    _p(doc,
        "Combined with our existing z-score-based modeling "
        "methodology — transparent, fan-customizable, in production "
        "today — the cross-sorted data bank gives us a foundation "
        "that no off-the-shelf data feed can match. The gambling "
        "product described in this document is what gets built on "
        "top of it."
    )

    # ════════════════════════════════════════════════════════════
    # 2. THE OPPORTUNITY
    # ════════════════════════════════════════════════════════════
    _h1(doc, "2. The Opportunity")

    _p(doc,
        "US sports gambling research is a growing market post-PASPA "
        "repeal, with multiple billion-dollar segments serving "
        "distinct customer profiles. Approaching this market well "
        "requires picking the right segment, not chasing all of them."
    )

    _h2(doc, "2.1 Market Segments and TAMs")

    _bold_label(doc, "Casual recreational bettors ($20–100/weekend, "
        "~30M Americans).",
        "Want quick takes — \"give me three picks for Sunday.\" "
        "High churn. Saturated by free Twitter content, podcasts, "
        "and freemium picks services. Willingness to pay $5–15/mo, "
        "but extracting that without massive marketing spend is "
        "hard. Not our target.")

    _bold_label(doc, "Serious recreational bettors ($500–2K/week, "
        "~2–4M Americans).",
        "Read research, follow sharp Twitter, track their own "
        "bets, line-shop across multiple books. Willingness to pay "
        "$30–100/mo for tools that produce real edge. Funded by "
        "their own betting profits, low churn when value is real. "
        "Big enough to support a single-founder product, too small "
        "to interest ESPN, The Athletic, or DraftKings' content "
        "arm. THIS IS OUR TARGET.")

    _bold_label(doc, "Professional sharps "
        "(<50,000 nationwide, $10K+ weekly).",
        "Build their own models, want APIs and raw data. "
        "Willingness to pay $100–500/mo for tooling. Demanding, "
        "expect white-glove support. Aspirational tier we'll capture "
        "incidentally; not a launch focus.")

    _h2(doc, "2.2 Why Serious Recreational Is the Right Pick")

    _p(doc,
        "Three reasons this segment is uniquely well-suited to what "
        "we've built:"
    )

    _bold_label(doc, "1. They value transparency.",
        "Recreational gamblers are skeptical of black-box \"models.\" "
        "Our entire methodology is z-score-based and fan-auditable. "
        "Every projection in our product can be inspected to its "
        "underlying components. This is a buying signal black-box "
        "services can't replicate.")

    _bold_label(doc, "2. They want education, not just verdicts.",
        "A $20/mo subscriber wants to understand WHY the lean is "
        "\"lean under\" — not just be told. Picks-services churn "
        "this audience constantly. Tools that teach retain them.")

    _bold_label(doc, "3. They're underserved by the existing market.",
        "Most paid gambling tools assume sharp-level expertise out "
        "of the gate. Their UI is jargon-heavy. Their charts assume "
        "you already know what an alt-line ladder is. Our Bet School "
        "Mode (described below) explicitly addresses this gap — a "
        "feature most competitors won't bother to build because their "
        "internal voice is \"real bettors don't need hand-holding.\"")

    _h2(doc, "2.3 Bet School Mode — the Brand Pillar")

    _p(doc,
        "The product runs in two modes, toggleable per user, applied "
        "across every screen and feature."
    )

    _bold_label(doc, "Standard Mode — for power users.",
        "Raw numerical outputs (\"+9.2% EV\", \"ρ = 0.71\", \"P10/"
        "P50/P90\"), dense tables, sharp jargon retained, faster "
        "scanning once fluent.")

    _bold_label(doc, "Bet School Mode — for newcomers and curious fans.",
        "Plain-English labels (Strong Lean / Skip / Fade), hover "
        "tooltips on every gambling term, auto-generated paragraphs "
        "explaining WHY a bet is +EV, \"What can go wrong?\" variance "
        "previews shown by default, Practice Mode with a fake $1,000 "
        "bankroll, bet diary that grades user decisions weekly.")

    _p(doc,
        "Both modes show the same data, the same edges, the same "
        "recommendations. The difference is the reading level."
    )

    _p(doc, "The strategic effect of the mode toggle:")
    _bullet(doc,
        "Expands the addressable market beyond seasoned bettors "
        "into the much larger fan audience.")
    _bullet(doc,
        "Removes the upgrade barrier that kills most gambling-tool "
        "signups in the first session.")
    _bullet(doc,
        "Practice Mode becomes a conversion engine: free → practice "
        "→ understand → pay.")
    _bullet(doc,
        "The brand pillar becomes durable — \"Lions Rater: smart "
        "enough for sharps, friendly enough for fans.\"")

    _h2(doc, "2.4 Pricing Tiers")

    table = doc.add_table(rows=4, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Tier"
    hdr[1].text = "Audience"
    hdr[2].text = "Price"
    hdr[3].text = "Active"
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    rows = [
        ("Free", "Anyone — the public face", "$0", "Year-round"),
        ("Standard Paid", "Serious recreational", "$15–20/mo",
         "In-season (Sept–Feb)"),
        ("Pro", "Sharps and power users", "$50–80/mo",
         "In-season (Sept–Feb)"),
    ]
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val

    _p(doc, "")

    _p(doc,
        "Free tier includes the full existing rater app, the player "
        "blurb engine, basic injury report viewing, default scenario "
        "projections, and a single-team daily digest. Standard Paid "
        "unlocks the full feature stack — six game-bet technologies, "
        "nine prop technologies, push notifications, the full "
        "scenario builder, PDF exports. Pro tier adds API access, "
        "earlier real-time alerts, and custom model-weight overrides."
    )

    _h3(doc, "Revenue math")
    _bullet(doc, "1,000 paid subs at $20/mo × 5 mo = $100K/year")
    _bullet(doc, "3,000 paid subs at blended $25/mo × 5 mo = $375K/year")
    _bullet(doc,
        "Single-founder break-even: ~100 paid subscribers")

    _p(doc,
        "The pricing is deliberate. $15–20/mo is "
        "recreational-accessible — \"what I'd spend on a parlay\" "
        "money. The product earns retention by producing real edge "
        "over the season. Reach is a primary goal alongside revenue: "
        "the free tier keeps the public face open, builds "
        "word-of-mouth, and serves as the conversion funnel into "
        "paid."
    )

    # ════════════════════════════════════════════════════════════
    # 3. COSTS
    # ════════════════════════════════════════════════════════════
    _h1(doc, "3. Costs")

    _p(doc,
        "All annual operating costs for the gambling product layer "
        "atop existing infrastructure."
    )

    cost_rows = [
        ("Category", "Vendor", "Cost", "Notes"),
        ("Live odds API (in-season)",
         "TheOddsAPI / OddsJam",
         "$30–60/mo × 5 mo",
         "Pause Mar–Aug"),
        ("Historical odds backfill (one-time)",
         "SBR Odds Portal",
         "$500–1,500",
         "NFL 2016–present"),
        ("News aggregation (in-season)",
         "NewsAPI + RSS",
         "$30/mo × 5 mo",
         "30-second lag"),
        ("Real-time news (optional v2)",
         "Twitter API v2",
         "$100–200/mo × 5 mo",
         "Sub-minute alerts"),
        ("Email service",
         "SendGrid",
         "$15–25/mo",
         "Digests + alerts"),
        ("Push notifications",
         "Pushover / Apple-Google",
         "<$10/mo",
         "Mobile alerts"),
        ("Weather forecast",
         "NOAA / OpenWeatherMap",
         "Free–$10/mo",
         "Free tier sufficient"),
        ("Hosting (current)",
         "Streamlit Cloud",
         "$0 → ~$50/mo at scale",
         "Currently free tier"),
        ("Payment processing",
         "Stripe",
         "2.9% + $0.30/txn",
         "~$3/sub/year"),
        ("Domain / branding",
         "Various",
         "~$50/yr",
         "Existing"),
    ]
    cost_table = doc.add_table(rows=len(cost_rows), cols=4)
    cost_table.style = "Light Grid Accent 1"
    for i, row in enumerate(cost_rows):
        for j, val in enumerate(row):
            cost_table.rows[i].cells[j].text = val
            if i == 0:
                for p in cost_table.rows[i].cells[j].paragraphs:
                    for r in p.runs:
                        r.bold = True

    _p(doc, "")
    _bold_label(doc, "Year 1 lean configuration:",
        "~$1,200–2,000. Without Twitter API, paid odds Sept–Feb only.")
    _bold_label(doc, "Year 1 premium configuration:",
        "~$2,500–4,000. With Twitter API for real-time news.")

    _h2(doc, "What We Are NOT Paying For")
    _bullet(doc, "nflverse play-by-play and roster data (free)")
    _bullet(doc, "CFBD college play-by-play and recruiting (free)")
    _bullet(doc,
        "NFL practice reports and injury archives (public via "
        "nflreadpy, free)")
    _bullet(doc, "Snap counts and PFR advanced stats (nflreadpy, free)")
    _bullet(doc,
        "The cross-indexed data bank itself — already built, three "
        "years of work, sunk cost")

    _p(doc,
        "Break-even at fewer than 100 paid subscribers in the lean "
        "configuration. The existing rater app already serves enough "
        "fan traffic for that scale to be conservatively achievable "
        "in season one."
    )

    # ════════════════════════════════════════════════════════════
    # 4. GAME BET TECHNOLOGIES
    # ════════════════════════════════════════════════════════════
    _h1(doc, "4. Game Bet Technologies")

    _p(doc,
        "Six features built on the cross-sorted data bank for sides, "
        "totals, and team-level betting research. Each uses "
        "dimensions of the data no off-the-shelf service has."
    )

    # Feature 4.1
    _feature_block(doc, "4.1 Injury Probability + Usage Retention Model",
        what=("For any player on a current week's injury report, "
              "output two numbers: the probability he plays Sunday, "
              "and — if he plays — his expected usage retention as "
              "a percentage of his season baseline. Backed by 5+ "
              "named historical comparable cases."),
        how=("Pulls nflreadpy.load_injuries() — the full historical "
             "archive of NFL practice reports back to 2009, "
             "structured with player_id, body part, Wed/Thu/Fri "
             "practice status (DNP/Limited/Full), and game-day "
             "designation. We normalize body parts into ~30 "
             "standardized buckets, then for any current case, find "
             "the historical (position, body_part, "
             "practice_sequence) cohort and compute the empirical "
             "probability of playing along with average usage "
             "retention if active."),
        sample=("Tua Tagovailoa, QB, shoulder. Practice this week: "
                "DNP-Wed / Limited-Thu / Limited-Fri. Cohort: 1,000+ "
                "historical cases. Probability of playing: 73%. If "
                "active, expects 84% of season-baseline pass yards. "
                "Comparables: A. Rodgers 2017 W12, P. Mahomes 2023 "
                "W16, R. Wilson 2019 W10."),
        diff=("Free sites show today's practice status. Paid sites "
              "show practice trend. Almost nobody computes a "
              "historical cohort base rate with named comparables. "
              "The named cases are the trust-builder for serious "
              "recreational bettors who don't trust black-box "
              "probabilities.")
    )

    # Feature 4.2
    _feature_block(doc,
        "4.2 Game-Script Simulator (with Cascading Multi-Injury)",
        what=("Toggle any starter — or any combination of starters — "
              "to OUT or LIMITED on the matchup card. Live updates "
              "the game script projection: predicted pass attempts, "
              "rush attempts, personnel grouping mix, backfield "
              "distribution, top-receiver target share, and final "
              "score model. Then re-prices every prop on the slate "
              "accordingly."),
        how=("From game_pbp.parquet, we bucket each (team, season) "
             "into \"starter played\" vs \"starter missed\" subsets "
             "and compute scheme deltas: pass rate, 11/12/21 "
             "personnel rates, shotgun rate, first-down pass rate, "
             "red-zone rush rate, play-action rate. A two-tier "
             "model handles the small-sample problem: a league-wide "
             "base rate (\"when teams lose RB1, league average pass "
             "rate +4%\") plus a team-specific overlay shrunk via "
             "Bayesian methods. For multi-injury scenarios, an "
             "interaction-term lookup table accounts for cascading "
             "effects (losing QB1 + WR1 together has a different "
             "impact than the sum of each individual loss)."),
        sample=("Toggle Gibbs to OUT. Pass rate +6%, 11 personnel "
                "+6%, first-down pass +9%. Predicted pass attempts: "
                "38 → 41. Pacheco projected 18 carries / 78 rush "
                "yards. WR1 target share: 9.2 → 10.1. Final score "
                "model: Lions 27 → 28."),
        diff=("Books shift the obvious props (the missing player's "
              "prop drops) but miss second-order effects: WR2 "
              "targets go up, defense's run defense matters less, "
              "the game total shifts. That gap is where edge lives, "
              "and nobody else is modeling it cleanly across all "
              "prop types simultaneously.")
    )

    # Feature 4.3
    _feature_block(doc,
        "4.3 Books-vs-Model Behavioral Baseline",
        what=("Tells the user not just \"the line moved\" but "
              "whether the move was consistent with how books "
              "historically react to that kind of news, or whether "
              "the book over- or under-reacted. Surfaces persistent "
              "edge patterns the market hasn't priced out."),
        how=("A historical odds archive (one-time backfill, "
             "~$500–1,500 from SBR Odds Portal for NFL game-level "
             "lines 2016–present) joined to the injury archive "
             "(Feature 4.1's data) and play-by-play outcomes. For "
             "each historical (news_type, body_part, position) "
             "tuple, we compute (1) the average book line move, "
             "(2) the actual game outcome relative to the original "
             "line, and (3) the gap between the two — i.e., the "
             "systematic over- or under-reaction."),
        sample=("Books historically over-move on big-name QB "
                "injuries (the public floods the under, the line "
                "overshoots — value going back to the over). Books "
                "under-react when the backup is a high-EPA quality "
                "player like Brock Purdy 2022 or Tyrod Taylor — "
                "the team's actual scoring barely moves but the "
                "line drops as if any backup is generic. Books "
                "reliably under-react to RB1 injuries because "
                "committees are common, but the actual top-backup "
                "gets 18+ touches and the prop is too low."),
        diff=("These patterns have been documented anecdotally in "
              "sharp-betting literature for years. Nobody currently "
              "publishes a clean public-facing tool that "
              "continuously surfaces them, because nobody bothered "
              "to backfill the data and run the model. "
              "First-season output: \"We backtested 8 seasons of "
              "NFL injury news against book line movement. Here are "
              "the 12 patterns where books reliably mis-react. Each "
              "week we tell you which of those patterns is live in "
              "the current slate.\"")
    )

    # Feature 4.4
    _feature_block(doc,
        "4.4 Smart Alerts + Digests",
        what=("Push notifications and email digests that bundle "
              "every other feature into a single, bet-actionable "
              "update. Competitors tell you the news. We tell you "
              "what the news means for Sunday's matchup."),
        how=("Source ingestion via NewsAPI + RSS scrapers from "
             "Pro Football Talk, ESPN, NFL.com, and team beat "
             "writers (~30-second lag). Detection filters incoming "
             "text for player names + status keywords, deduplicates "
             "across sources. Each detected event is matched "
             "against subscribers' watchlists and tracked bets, "
             "then a single notification fuses Feature 4.1's "
             "probability, Feature 4.2's game-script reshape, the "
             "prop re-pricing, the historical book-behavior context "
             "(Feature 4.3), and the user's affected open positions."),
        sample=("ALERT: Jared Goff officially out vs Texans "
                "(shoulder). Hooker historical retention 71% of "
                "starter baseline (cohort 800 backup-QB games, HIGH "
                "confidence). Lions team total 23.5 → 22.0. Pass "
                "rate 58% → 64%. Hooker pass yards 245.5 line → "
                "model 218 (UNDER, 27.5 yds edge). Pacheco rush "
                "yards 78.5 line → model 92 (OVER, 13.5 yds edge). "
                "Books historically under-react to mid-week "
                "starting-QB news by ~11 yards; current line move "
                "is 8 yards — edge persists. Your Goff over 244.5 "
                "is invalidated; available pivot: Pacheco OVER 78.5."),
        diff=("Competitors stop at the headline. Even premium "
              "services don't fuse all four prior features into "
              "one notification. We can — because all features "
              "sit on the same data foundation. Every alert is 30 "
              "minutes of manual research delivered as a single "
              "push.")
    )

    # Feature 4.5
    _feature_block(doc,
        "4.5 Weather Production Window",
        what=("Real-time slider-driven weather modeling. For any "
              "player on the slate, the user adjusts temperature, "
              "wind speed, precipitation, and surface, and watches "
              "the player's projected production range "
              "(P10/P50/P90) update live. Cohort-matched to the "
              "player's historical games in similar conditions, "
              "with sample-size confidence flags."),
        how=("Every play in game_pbp.parquet carries weather "
             "observations. For any input weather scenario, we "
             "query each player's historical games matching the "
             "conditions within tolerance bounds, compute the "
             "empirical distribution of stat outcomes, and surface "
             "the 10th/50th/90th percentiles. When a player's "
             "individual sample is thin (<5 games), the system "
             "blends in a similar-tier-player cohort with the "
             "player's baseline as the anchor. Confidence is "
             "labeled visibly: HIGH (15+ games), MEDIUM (6–14), "
             "LOW (<5 with cohort fallback). Wind direction × "
             "stadium azimuth is also computed (a 20mph crosswind "
             "ruins deep balls; a 20mph tailwind helps them)."),
        sample=("Goff at 36°F, 14 mph wind, snow, grass. Cohort: 7 "
                "historical games. P10–P50–P90: 187 / 221 / 256. "
                "Confidence: MEDIUM. User drags temperature to 43°F, "
                "sun: cohort expands to 22 games, P10–P50–P90 shifts "
                "to 214 / 256 / 291. Confidence: HIGH."),
        diff=("Competitor weather analysis is binary (\"wind > 15 "
              "mph = take under\"). Mid-tier services apply static "
              "league-average wind impact. Lions Rater is the only "
              "product that produces a continuous, player-specific "
              "production range that morphs in real time as the "
              "user explores conditions, with honest confidence "
              "labels.")
    )

    # Feature 4.6 — NEW
    _feature_block(doc,
        "4.6 Coaching & Scheme Tendency Layer",
        what=("Models coaching staff preferences from career "
              "play-by-play: pass/run rate by down and distance, "
              "personnel mix, red-zone aggression, 4th-down "
              "aggression, 2-minute drill efficiency, play-action "
              "rate, blitz rate, disguise rate. Feeds team-level "
              "projections — pace, total points, situational "
              "scoring — and surfaces coach-on-coach matchup "
              "history when relevant."),
        how=("Cross-indexed pbp tagged by play-caller (OC) and "
             "defensive coordinator. We compute career tendencies "
             "across each coach's tenures so the model carries "
             "forward correctly when an OC changes teams (Ben "
             "Johnson moving jobs, McVay's offensive DNA "
             "traveling with his coordinators). Tendencies are "
             "z-scored against league averages for direct "
             "comparison and used as adjustment multipliers on "
             "team-level projections."),
        sample=("Ben Johnson runs play-action on 32% of dropbacks "
                "vs league average 22%. Texans defense ranks 24th "
                "vs play-action by EPA. Adjusts Lions team total "
                "+1.5 vs baseline."),
        diff=("Most sites show team tendencies, which conflate "
              "coach + roster signal. We isolate the play-caller "
              "signal across coaching tenures. When a coach moves "
              "teams, our model travels with him — competitor "
              "models reset to team averages.")
    )

    # ════════════════════════════════════════════════════════════
    # 5. PROP BET TECHNOLOGIES
    # ════════════════════════════════════════════════════════════
    _h1(doc, "5. Prop Bet Technologies")

    _p(doc,
        "Nine features for player-prop research. Player props are "
        "the highest-margin gambling segment for an indie product "
        "because the markets are smaller, the books move slower, "
        "and our cross-sorted data has the granular per-player "
        "per-stat splits books haven't fully priced in."
    )

    _feature_block(doc, "5.1 Decomposed Prop Projection",
        what=("Every prop projection ships with a line-by-line "
              "breakdown showing exactly where each yard came "
              "from — injuries, splits, matchups, weather, scheme "
              "changes — auditable instead of black-box."),
        how=("The existing rater methodology — z-scored stats, "
             "bundle weights, formation/personnel splits — already "
             "produces these decompositions internally. We surface "
             "them. Each row is a stat z-col, situational split, "
             "or matchup factor the user can audit."),
        sample=("Tua pass yards — book line 235, our model 219, "
                "edge 4 yards under. Where the 26 yards of model "
                "gap come from: -8 shoulder injury cohort retention "
                "(Feature 4.1), -7 Hill out → top-receiver target "
                "share shift (Feature 4.2), -5 cold weather road "
                "game (Feature 4.5), -3 Texans top-5 in "
                "completion-pct allowed, -3 game script projects "
                "fewer dropbacks."),
        diff=("Most paid services say \"lean under\" and stop. "
              "Black box. The transparency itself is the product — "
              "serious recreational bettors are skeptical of "
              "models they can't audit.")
    )

    _feature_block(doc, "5.2 Same-Game Parlay (SGP) Correlation Edge",
        what=("Books price same-game parlays as if the legs are "
              "independent events. They aren't — props in the "
              "same game flow from the same plays and correlate "
              "heavily. We compute actual correlation from "
              "play-by-play and surface SGPs where the joint "
              "probability beats book pricing."),
        how=("For every player pair and every (player, team-stat) "
             "combination within the same offense, we pre-compute "
             "correlation matrices from historical games. When "
             "the user builds an SGP, we look up the correlation "
             "between each leg pair, solve the joint probability "
             "with a multivariate-Gaussian or empirical-distribution "
             "method, and compare against the book's "
             "independent-multiplication price."),
        sample=("Three-leg SGP: Goff over 244.5 + St. Brown over "
                "75.5 + Lions team total over 24.5. Book pricing "
                "(independent): 13% → +650. Our pricing "
                "(correlated): 32% → +210. Edge: +440 yards of "
                "value [STRONG LEAN]."),
        diff=("The edge has been documented in sharp-betting "
              "literature for years; no existing public tool "
              "handles it cleanly across user-built SGPs. We make "
              "it accessible to recreational bettors who couldn't "
              "compute joint probabilities manually.")
    )

    _feature_block(doc,
        "5.3 Alt-Line EV Finder (with Ladder for Dummies)",
        what=("For every prop, we compute the model's projected "
              "probability at every rung of the alt-line ladder. "
              "The system surfaces the rung with the best expected "
              "value — which is rarely the main line."),
        how=("Each prop's projection is a full distribution "
             "(mean + standard deviation + skew from historical "
             "games), not a point estimate. For each alt-line "
             "threshold and price, we compute P(yards > threshold) "
             "from our distribution, then EV = P × (payout − 1) − "
             "(1 − P) × 1. Sort by EV, surface the best rung. In "
             "Bet School Mode, this entire feature includes "
             "step-by-step explanations: what an alt-line ladder "
             "is, how American odds work, why this rung is +EV, "
             "and what the variance looks like over 100 such bets."),
        sample=("Goff pass yards. Main line over 245.5 at -110 "
                "(model: 51%, EV -2.6% — fade). Alt over 274.5 at "
                "+180 (model: 39%, EV +9.2% — best EV). Alt over "
                "285.5 at +250 (model: 32%, EV +12% — big value)."),
        diff=("Most bettors leave 5–15% of EV on the table by "
              "reflexively betting main lines. This feature finds "
              "the rung where the book is most wrong. Combined "
              "with Bet School, it's also a teaching tool for "
              "newcomers learning the alt-line concept.")
    )

    _feature_block(doc, "5.4 Smart Parlay Builder",
        what=("A daily ranked menu of mispriced props is the "
              "starting point. The user builds any parlay by "
              "checking boxes; the tool computes joint probability "
              "with full correlation handling, compares against "
              "the book's quoted parlay odds, and actively "
              "suggests improvements: \"drop leg X, add leg Y, "
              "total EV jumps.\""),
        how=("Re-uses Feature 5.1 (decomposed projections) for the "
             "menu and Feature 5.2 (correlation matrices) for the "
             "joint math. The suggestion engine evaluates each "
             "candidate next-leg's marginal contribution to joint "
             "EV. Conflict detection flags anti-correlated legs. "
             "Auto-optimized parlays (\"show me the best 3-leg / "
             "4-leg / 5-leg parlays today\"), risk-tolerance modes "
             "(steady-value vs lottery-ticket), a hedge calculator, "
             "and a personal track record system are sub-features."),
        sample=("User selects three legs; tool detects two "
                "anti-correlate, suggests dropping one, adds an "
                "alternative that correlates positively with the "
                "remaining legs. Parlay moves from book +650 / our "
                "+210 to book +480 / our +98. EV improves "
                "dramatically."),
        diff=("Once a recreational gambler uses this for a few "
              "weekends, they won't go back to building parlays "
              "manually. It's the stickiest feature on the prop "
              "side.")
    )

    _feature_block(doc,
        "5.5 Anytime / First TD Probability Vector",
        what=("Models touchdown probability per player per game, "
              "combining red-zone snap share, target share inside "
              "the 10, opponent goal-line defense rank, and "
              "projected game script. Decomposes anytime-TD into "
              "rushing-TD and receiving-TD probabilities so users "
              "can find which sub-prop is mispriced."),
        how=("Cross-indexed pbp gives us RZ snap participation, "
             "RZ target share, and goal-line defensive splits. "
             "We combine these with game-script likelihood (from "
             "spread + total) to produce a per-player TD "
             "probability decomposed by score type."),
        sample=("Pacheco anytime TD +110 (model says fair). "
                "Rushing TD only at +180 (model says 47% likely, "
                "fair price -113). Strong lean on rushing TD only."),
        diff=("Books are notoriously slow on TD markets because "
              "the math is hard for casuals — low base rates, "
              "multiple interacting factors. Our cross-sorted "
              "data computes RZ snap share, goal-line defense "
              "splits, and game-script likelihood all in one "
              "pipeline.")
    )

    _feature_block(doc,
        "5.6 Snap-Share / Target-Share Trend Divergence",
        what=("Flags props where a player's last-3-week usage "
              "(snap share, target share, RZ touches) diverges "
              "from the season baseline. Books often anchor to "
              "season averages and miss recent role expansions or "
              "contractions. Includes special handling for rookies "
              "and recent breakouts whose season prior is "
              "unreliable."),
        how=("For every offensive skill player, we maintain "
             "rolling 3-week and season-long usage metrics. Any "
             "stat z-score gap >0.5σ between recent and season "
             "triggers a divergence flag. Adjusts the prop "
             "projection toward recent form when the divergence "
             "is large enough to be statistically meaningful."),
        sample=("Loveland rec yards prop 32.5 — season avg 4.2 "
                "receptions, but last 3 weeks 7.3 receptions. Role "
                "expansion not yet priced in. Lean OVER."),
        diff=("Most sites show the season trend. We flag the "
              "specific moments when current-form has decoupled "
              "from season averages, which is exactly when the "
              "prop is most likely mispriced.")
    )

    _feature_block(doc, "5.7 Longest-Play Edge Finder",
        what=("Uses our explosive-rate z-scores per player to "
              "find mispriced \"longest rush\" and \"longest "
              "reception\" props. Players with elite explosive "
              "rates have structurally undervalued longest-play "
              "markets because books model these on smooth "
              "distributions while reality is bimodal."),
        how=("Per-player explosive-rate (yards-per-play above the "
             "75th percentile of all plays) is already in our z-col "
             "stack. Combined with empirical distribution of the "
             "player's explosive plays, we model the probability "
             "of any single play exceeding the threshold X over a "
             "given number of opportunities (carries or targets)."),
        sample=("Achane explosive run rate +2.4σ. \"Longest rush "
                "over 25 yards\" prop priced at +110. Model says "
                "64% likely. EV: +21%."),
        diff=("Almost no public service models longest-play markets "
              "this directly. The data is sitting in our pbp.")
    )

    _feature_block(doc,
        "5.8 Defense-vs-Position Stat Allowance Tables",
        what=("Every NFL defense ranked by yards / TDs / receptions "
              "/ INTs allowed per opposing position group, broken "
              "down by stat type. Browseable, sortable, used as "
              "the foundational matchup-research layer that feeds "
              "the projections in Features 5.1, 5.5, 5.6."),
        how=("Aggregate every defensive snap from game_pbp.parquet "
             "by opposing position group, compute z-scored "
             "allowance per stat per defense per season. Surface "
             "as a sortable table with filters."),
        sample=("Texans defense: 32nd vs slot WRs in 11 personnel. "
                "St. Brown matchup advantage: significant. Adjusts "
                "St. Brown projection +9 yds vs season baseline."),
        diff=("This is table-stakes for prop research — every paid "
              "prop service leads with it. Without it, sharps will "
              "assume our product is missing core functionality. "
              "With it, we cover the foundational ground while "
              "still differentiating with the eight other features.")
    )

    # Feature 5.9 — NEW
    _feature_block(doc,
        "5.9 Coaching Tendency Adjustment",
        what=("Per-player prop projections adjusted for the "
              "play-caller's specific tendencies. Pass-heavy OCs "
              "inflate QB and WR yardage props. Run-heavy OCs "
              "inflate RB carries and yardage props. Screen-heavy "
              "schemes (Mike McDaniel, Kyle Shanahan) inflate RB "
              "receptions and slot-WR yards-after-catch. "
              "Disguise-heavy DCs inflate QB sack and INT props."),
        how=("Each coach's tendency profile (from Feature 4.6's "
             "engine) is applied as a per-stat multiplier on every "
             "affected player's projection. The OC's career "
             "screen-rate, deep-ball rate, RZ rush rate, etc. all "
             "feed forward to the player props of the players he "
             "currently calls plays for."),
        sample=("Lions use screens on 11% of plays (league avg "
                "4%). St. Brown's screen-yards-after-catch prop "
                "is structurally undervalued — adjusts +6 yards "
                "on his rec-yards line."),
        diff=("Most paid prop services apply team-level adjustments. "
              "We adjust at the OC/DC level, which is the right "
              "granularity — and our data lets us track tendencies "
              "across coach moves rather than tying them to "
              "franchise.")
    )

    # ════════════════════════════════════════════════════════════
    # 6. Questions for the Sharp
    # ════════════════════════════════════════════════════════════
    _h1(doc, "6. Questions for the Sharp")

    _p(doc,
        "Open questions for our discussion."
    )

    _bullet(doc,
        "Which of these 15 features hits hardest? Which would you "
        "actually pay $20–80/mo for?")
    _bullet(doc,
        "What's missing from this stack that you currently pay for "
        "elsewhere?")
    _bullet(doc,
        "What patterns of book mis-reaction have you seen that we "
        "haven't listed?")
    _bullet(doc,
        "Is the named-comparable-cases approach (e.g., \"comp to "
        "Rodgers 2017 W12\") trust-building, or noise?")
    _bullet(doc,
        "If we had to cut three features for a v1 launch, which "
        "three?")
    _bullet(doc,
        "Honest read on the price points — too high, too low, "
        "just right?")
    _bullet(doc,
        "Would you ever switch books based on our alt-line EV "
        "finder, or is friction too high?")
    _bullet(doc,
        "How do you feel about the Bet School / Standard Mode "
        "split — does it dilute the product for power users, or "
        "is the larger TAM worth it?")

    # ════════════════════════════════════════════════════════════
    # 7. Appendix
    # ════════════════════════════════════════════════════════════
    _h1(doc, "7. Appendix: Live App")

    _p(doc, "Production app: lions-rater.streamlit.app")
    _p(doc,
        "The current rater app is fan-facing only (no gambling "
        "features). The 15-feature stack above sits on top of the "
        "same data foundation, gated behind a paid tier, active "
        "during the season."
    )

    doc.save(OUT)
    print(f"✓ wrote {OUT}")


if __name__ == "__main__":
    build()
